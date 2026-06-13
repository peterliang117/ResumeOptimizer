#!/usr/bin/env python3
"""Generate truthful resume tailoring suggestions and optionally apply accepted edits."""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, quote, urlparse


SUGGESTION_SCHEMA: dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "job_summary": {"type": "string"},
        "must_have_skills": {"type": "array", "items": {"type": "string"}},
        "nice_to_have_skills": {"type": "array", "items": {"type": "string"}},
        "matched_evidence": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "requirement": {"type": "string"},
                    "evidence": {"type": "string"},
                    "confidence": {"type": "string", "enum": ["high", "medium", "low"]},
                },
                "required": ["requirement", "evidence", "confidence"],
            },
        },
        "suggested_edits": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "paragraph_id": {"type": "string"},
                    "original": {"type": "string"},
                    "suggested": {"type": "string"},
                    "reason": {"type": "string"},
                    "evidence_source": {
                        "type": "string",
                        "enum": ["resume", "profile", "resume_and_profile"],
                    },
                    "truth_risk": {"type": "string", "enum": ["low", "medium", "high"]},
                },
                "required": [
                    "paragraph_id",
                    "original",
                    "suggested",
                    "reason",
                    "evidence_source",
                    "truth_risk",
                ],
            },
        },
        "confirmation_questions": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "requirement": {"type": "string"},
                    "question": {"type": "string"},
                },
                "required": ["requirement", "question"],
            },
        },
    },
    "required": [
        "job_summary",
        "must_have_skills",
        "nice_to_have_skills",
        "matched_evidence",
        "suggested_edits",
        "confirmation_questions",
    ],
}


class LLMUnavailableError(RuntimeError):
    """Raised when an optional LLM backend cannot be used."""


@dataclass(frozen=True)
class ResumeParagraph:
    paragraph_id: str
    text: str


def iter_document_paragraphs(document: Any) -> list[tuple[str, Any]]:
    paragraphs: list[tuple[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        paragraphs.append((f"body:p{index}", paragraph))

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    paragraphs.append(
                        (
                            f"table:{table_index}:r{row_index}:c{cell_index}:p{paragraph_index}",
                            paragraph,
                        )
                    )
    return paragraphs


def load_job_text(job: str | None = None, job_url: str | None = None) -> str:
    if job_url:
        return load_url_text(job_url)
    if not job:
        raise SystemExit("Provide either --job or --job-url.")
    if re.match(r"^https?://", job):
        return load_url_text(job)

    path = Path(job)
    if path.exists():
        return path.read_text(encoding="utf-8")
    return job


def load_url_text(url: str) -> str:
    if not re.match(r"^https?://", url):
        raise SystemExit("--job-url must start with http:// or https://")

    try:
        import requests
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise SystemExit("Install requests and beautifulsoup4 to load job URLs.") from exc

    parsed_url = urlparse(url)
    if parsed_url.netloc.lower() == "jobs.ashbyhq.com":
        path_parts = [part for part in parsed_url.path.split("/") if part]
        if len(path_parts) >= 2:
            organization, job_id = path_parts[:2]
            api_url = f"https://api.ashbyhq.com/posting-api/job-board/{organization}"
            try:
                api_response = requests.get(api_url, timeout=20)
                api_response.raise_for_status()
                jobs = api_response.json().get("jobs", [])
            except (requests.RequestException, ValueError) as exc:
                raise SystemExit(f"Could not fetch Ashby job board API: {exc}") from exc

            for posting in jobs:
                posting_url = str(posting.get("jobUrl", ""))
                posting_id = urlparse(posting_url).path.rstrip("/").split("/")[-1]
                if posting_id != job_id:
                    continue
                description = str(posting.get("descriptionPlain", "")).strip()
                if not description:
                    break
                header = "\n".join(
                    value
                    for value in [
                        str(posting.get("title", "")).strip(),
                        str(posting.get("location", "")).strip(),
                        str(posting.get("workplaceType", "")).strip(),
                        str(posting.get("employmentType", "")).strip(),
                    ]
                    if value
                )
                return f"{header}\n\n{description}".strip()

            raise SystemExit(
                "Ashby job was not found on the employer's current public board. "
                "The posting may be unlisted or expired."
            )

    try:
        response = requests.get(
            url,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/124.0 Safari/537.36"
                )
            },
            timeout=20,
        )
        response.raise_for_status()
    except requests.RequestException as exc:
        raise SystemExit(
            f"Could not fetch job URL: {exc}\n"
            "If the job board blocks automated access, paste the description into jobs/job.txt "
            "and use --job jobs/job.txt."
        ) from exc

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()
    text = soup.get_text("\n")
    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(cleaned) < 200:
        raise SystemExit(
            "Fetched page contained very little readable text. "
            "The job board may require JavaScript or login; paste the description into jobs/job.txt."
        )
    return cleaned


def load_resume_paragraphs(path: Path) -> list[ResumeParagraph]:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("Install python-docx to read DOCX resumes.") from exc

    document = Document(path)
    paragraphs: list[ResumeParagraph] = []
    for paragraph_id, paragraph in iter_document_paragraphs(document):
        text = paragraph.text.strip()
        if text:
            paragraphs.append(ResumeParagraph(paragraph_id=paragraph_id, text=text))
    return paragraphs


def build_prompt(
    resume_paragraphs: list[ResumeParagraph],
    job_text: str,
    profile_facts: str,
) -> str:
    resume_json = [{"paragraph_id": p.paragraph_id, "text": p.text} for p in resume_paragraphs]
    return f"""
You are helping tailor a one-page resume to a job description.

Hard rules:
- Do not invent or exaggerate facts.
- Only suggest changes supported by the resume or profile facts.
- Prefer compact edits that keep the resume one page.
- Do not add unsupported tools, metrics, employers, dates, degrees, or responsibilities.
- If a job requirement is not supported, ask a confirmation question instead of adding it.
- Suggest edits to existing paragraphs only, using paragraph_id.

Resume paragraphs:
{json.dumps(resume_json, indent=2)}

Profile facts allowed as evidence:
{profile_facts}

Job description:
{job_text}
""".strip()


def normalize_azure_endpoint_candidate(value: str) -> str:
    candidate = value.strip()
    if not candidate:
        return ""
    if candidate.lower().startswith("https") and "://" not in candidate:
        candidate = "https://" + candidate[5:]
    marker = ".cognitiveservices.azure.com"
    marker_index = candidate.lower().find(marker)
    if marker_index != -1:
        return candidate[: marker_index + len(marker)]
    return candidate.rstrip("/")


def discover_azure_keys_file() -> Path | None:
    configured = os.getenv("AZURE_OPENAI_KEYS_FILE", "").strip()
    if configured:
        return Path(configured)
    for base in [Path.cwd(), *Path.cwd().parents]:
        candidate = base / "keys.txt"
        if candidate.exists():
            return candidate
    return None


def parse_azure_keys_lines(lines: list[str]) -> dict[str, str]:
    payload: dict[str, str] = {}
    positional: list[str] = []
    key_map = {
        "AZURE_OPENAI_ENDPOINT": "endpoint",
        "ENDPOINT": "endpoint",
        "AZURE_OPENAI_API_KEY": "api_key",
        "API_KEY": "api_key",
        "KEY": "api_key",
        "AZURE_OPENAI_DEPLOYMENT": "deployment",
        "DEPLOYMENT": "deployment",
        "AZURE_OPENAI_API_VERSION": "api_version",
        "API_VERSION": "api_version",
    }
    for line in lines:
        if "=" in line:
            name, value = line.split("=", 1)
            mapped = key_map.get(name.strip().upper())
            if mapped:
                payload[mapped] = value.strip().strip('"').strip("'")
                continue
        positional.append(line)

    if positional:
        payload.setdefault("endpoint", positional[0])
    if len(positional) > 1:
        payload.setdefault("api_key", positional[1])
    if len(positional) > 2:
        payload.setdefault("deployment", positional[2])
    return payload


def read_azure_keys_file() -> dict[str, str]:
    path = discover_azure_keys_file()
    if not path:
        return {}
    if not path.exists():
        raise LLMUnavailableError(f"AZURE_OPENAI_KEYS_FILE does not exist: {path}")
    lines = [
        line.strip().lstrip("\ufeff")
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]
    if len(lines) < 2:
        raise LLMUnavailableError("AZURE_OPENAI_KEYS_FILE must contain an endpoint line and an API key line.")
    payload = parse_azure_keys_lines(lines)
    if len(payload.get("api_key", "")) < 20:
        raise LLMUnavailableError("AZURE_OPENAI_KEYS_FILE API key line is missing or too short.")
    endpoint = normalize_azure_endpoint_candidate(payload.get("endpoint", ""))
    parsed = urlparse(endpoint)
    if parsed.scheme != "https" or not parsed.netloc:
        raise LLMUnavailableError("AZURE_OPENAI_KEYS_FILE endpoint line is not a valid Azure endpoint.")
    api_version = ""
    original = payload.get("endpoint", "")
    if "://" in original:
        api_version = parse_qs(urlparse(original).query).get("api-version", [""])[0]
    return {
        "endpoint": endpoint,
        "api_key": payload.get("api_key", ""),
        "deployment": payload.get("deployment", ""),
        "api_version": payload.get("api_version", "") or api_version,
    }


def has_azure_config() -> bool:
    try:
        keys_payload = read_azure_keys_file()
    except LLMUnavailableError:
        return False
    return bool(
        (os.getenv("AZURE_OPENAI_ENDPOINT", "").strip() or keys_payload.get("endpoint"))
        and (os.getenv("AZURE_OPENAI_DEPLOYMENT", "").strip() or keys_payload.get("deployment"))
        and (
            os.getenv("AZURE_OPENAI_API_KEY", "").strip()
            or keys_payload.get("api_key")
            or Path(os.getenv("AZURE_OPENAI_API_KEY_PATH", ".secrets/azure_openai_api_key")).exists()
        )
    )


def read_azure_api_key(keys_payload: dict[str, str] | None = None) -> str:
    if keys_payload and keys_payload.get("api_key"):
        return keys_payload["api_key"]

    api_key = os.getenv("AZURE_OPENAI_API_KEY", "").strip()
    if api_key:
        return api_key

    key_path = Path(os.getenv("AZURE_OPENAI_API_KEY_PATH", ".secrets/azure_openai_api_key"))
    if not key_path.exists():
        raise LLMUnavailableError(
            "Azure OpenAI API key not found. Set AZURE_OPENAI_API_KEY or "
            f"create local-only key file {key_path}."
        )
    api_key = key_path.read_text(encoding="utf-8").strip()
    if not api_key:
        raise LLMUnavailableError(f"Azure OpenAI API key file is empty: {key_path}")
    return api_key


def analyze_with_azure_openai(prompt: str) -> dict[str, Any]:
    try:
        import requests
    except ImportError as exc:
        raise LLMUnavailableError("Install requests to use Azure OpenAI analysis.") from exc

    keys_payload = read_azure_keys_file()
    endpoint = (os.getenv("AZURE_OPENAI_ENDPOINT", "").strip() or keys_payload.get("endpoint", "")).rstrip("/")
    deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT", "").strip() or keys_payload.get("deployment", "")
    api_version = (
        os.getenv("AZURE_OPENAI_API_VERSION", "").strip()
        or keys_payload.get("api_version", "")
        or "2025-04-01-preview"
    )
    if not endpoint or not deployment:
        raise LLMUnavailableError("AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_DEPLOYMENT are required.")

    api_key = read_azure_api_key(keys_payload)
    deployment_path = quote(deployment, safe="")
    url = f"{endpoint}/openai/deployments/{deployment_path}/chat/completions?api-version={api_version}"
    payload = {
        "messages": [
            {
                "role": "system",
                "content": "Return only structured JSON. Be conservative and truth-preserving.",
            },
            {"role": "user", "content": prompt},
        ],
        "max_completion_tokens": 1800,
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": "resume_tailoring_suggestions",
                "schema": SUGGESTION_SCHEMA,
                "strict": True,
            },
        },
    }
    response = requests.post(
        url,
        headers={"Content-Type": "application/json", "api-key": api_key},
        json=payload,
        timeout=90,
    )
    if response.status_code >= 400:
        raise LLMUnavailableError(f"Azure OpenAI API error {response.status_code}: {response.text}")

    data = response.json()
    try:
        content = data["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as exc:
        raise LLMUnavailableError("Azure OpenAI response did not include message content.") from exc
    if not isinstance(content, str) or not content.strip():
        raise LLMUnavailableError("Azure OpenAI response content was empty.")
    return parse_json_text(content, source="Azure OpenAI")


def parse_json_text(text: str, *, source: str) -> dict[str, Any]:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        payload = json.loads(text)
    except json.JSONDecodeError as exc:
        raise LLMUnavailableError(f"{source} response was not valid JSON: {exc}") from exc
    if not isinstance(payload, dict):
        raise LLMUnavailableError(f"{source} response JSON must be an object.")
    return payload


def extract_output_text(response_payload: dict[str, Any]) -> str:
    if isinstance(response_payload.get("output_text"), str):
        return response_payload["output_text"]

    chunks: list[str] = []
    for item in response_payload.get("output", []):
        for content in item.get("content", []):
            text = content.get("text")
            if isinstance(text, str):
                chunks.append(text)
    if not chunks:
        raise SystemExit("OpenAI response did not include output text.")
    return "".join(chunks)


def fallback_analysis(
    resume_paragraphs: list[ResumeParagraph],
    job_text: str,
    reason: str | None = None,
) -> dict[str, Any]:
    stopwords = {
        "and",
        "are",
        "base",
        "build",
        "company",
        "days",
        "for",
        "from",
        "have",
        "into",
        "looking",
        "needs",
        "our",
        "per",
        "posted",
        "range",
        "requires",
        "responsibilities",
        "role",
        "salary",
        "senior",
        "that",
        "the",
        "this",
        "through",
        "with",
        "will",
        "year",
        "you",
        "your",
    }
    tokens = sorted(
        {
            token.strip(".,:;()[]{}").lower()
            for token in re.findall(r"[A-Za-z][A-Za-z+#.\-]{2,}", job_text)
            if token.strip(".,:;()[]{}").lower() not in stopwords
        }
    )
    resume_text = "\n".join(p.text for p in resume_paragraphs).lower()
    matched = [token for token in tokens if token in resume_text][:30]
    missing = [token for token in tokens if token not in resume_text][:20]
    return {
        "job_summary": (
            "Fallback keyword scan. Use Codex/manual review for rewrite suggestions."
            if not reason
            else f"Fallback keyword scan. Use Codex/manual review for rewrite suggestions. Reason: {reason}"
        ),
        "must_have_skills": matched,
        "nice_to_have_skills": [],
        "matched_evidence": [
            {"requirement": token, "evidence": "Keyword appears in resume.", "confidence": "medium"}
            for token in matched
        ],
        "suggested_edits": [],
        "confirmation_questions": [
            {"requirement": token, "question": f"Can you truthfully claim experience with {token}?"}
            for token in missing
        ],
    }


def analyze_resume_tailoring(
    resume_paragraphs: list[ResumeParagraph],
    job_text: str,
    profile_facts: str,
    *,
    provider: str,
    model: str,
) -> dict[str, Any]:
    prompt = build_prompt(resume_paragraphs, job_text, profile_facts)
    if provider == "none":
        return fallback_analysis(resume_paragraphs, job_text, "LLM provider disabled.")

    if provider == "azure" or (provider == "auto" and has_azure_config()):
        try:
            suggestions = analyze_with_azure_openai(prompt)
            suggestions["_analysis_backend"] = "azure_openai"
            return suggestions
        except LLMUnavailableError as exc:
            return fallback_analysis(resume_paragraphs, job_text, f"Azure OpenAI unavailable: {exc}")

    return fallback_analysis(resume_paragraphs, job_text, "Azure OpenAI is not configured.")


def validate_edit(edit: dict[str, Any], paragraph_by_id: dict[str, str]) -> str | None:
    paragraph_id = edit.get("paragraph_id")
    original = edit.get("original", "")
    suggested = edit.get("suggested", "")
    if paragraph_id not in paragraph_by_id:
        return f"Unknown paragraph_id: {paragraph_id}"
    if paragraph_by_id[paragraph_id].strip() != original.strip():
        return f"Original text mismatch for {paragraph_id}"
    if not suggested.strip():
        return f"Empty suggestion for {paragraph_id}"
    if edit.get("truth_risk") == "high":
        return f"Refusing high truth-risk edit for {paragraph_id}"
    return None


def apply_edits(source_docx: Path, output_docx: Path, accepted_edits: list[dict[str, Any]]) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("Install python-docx to edit DOCX resumes.") from exc

    document = Document(source_docx)
    paragraph_text = {
        paragraph_id: paragraph.text.strip()
        for paragraph_id, paragraph in iter_document_paragraphs(document)
        if paragraph.text.strip()
    }

    for edit in accepted_edits:
        error = validate_edit(edit, paragraph_text)
        if error:
            raise SystemExit(error)

    edits_by_id = {edit["paragraph_id"]: edit for edit in accepted_edits}
    for paragraph_id, paragraph in iter_document_paragraphs(document):
        if paragraph_id in edits_by_id:
            replace_paragraph_text_keep_style(paragraph, edits_by_id[paragraph_id]["suggested"])

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def replace_paragraph_text_keep_style(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Tailor a DOCX resume conservatively.")
    parser.add_argument("--resume", required=True, type=Path, help="Source DOCX resume.")
    job_input = parser.add_mutually_exclusive_group(required=True)
    job_input.add_argument("--job", help="Job text file or raw pasted job text.")
    job_input.add_argument("--job-url", help="URL of the job post to fetch and analyze.")
    parser.add_argument("--profile", type=Path, default=Path("profile/facts.md"))
    parser.add_argument("--out", required=True, type=Path, help="Output DOCX path.")
    parser.add_argument("--suggestions-out", type=Path, default=Path("outputs/suggestions.json"))
    parser.add_argument("--accepted-edits", type=Path, help="JSON file with accepted suggested_edits.")
    parser.add_argument("--dry-run", action="store_true", help="Only write suggestions JSON.")
    parser.add_argument("--model", default=None)
    parser.add_argument(
        "--llm-provider",
        choices=["auto", "azure", "none"],
        default=os.getenv("RESUME_OPTIMIZER_LLM_PROVIDER", "auto"),
        help=(
            "Suggestion backend. auto uses Azure OpenAI when AZURE_OPENAI_* is configured, "
            "then falls back to local/Codex manual review."
        ),
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if not args.resume.exists():
        raise SystemExit(f"Resume not found: {args.resume}")

    job_text = load_job_text(job=args.job, job_url=args.job_url)
    profile_facts = args.profile.read_text(encoding="utf-8") if args.profile.exists() else ""
    resume_paragraphs = load_resume_paragraphs(args.resume)

    if args.dry_run or not args.accepted_edits:
        suggestions = analyze_resume_tailoring(
            resume_paragraphs,
            job_text,
            profile_facts,
            provider=args.llm_provider,
            model=args.model,
        )
        write_json(args.suggestions_out, suggestions)
        print(f"Wrote suggestions: {args.suggestions_out}")
        if args.dry_run:
            return 0

    if args.accepted_edits:
        accepted_payload = json.loads(args.accepted_edits.read_text(encoding="utf-8"))
        accepted_edits = (
            accepted_payload.get("suggested_edits", accepted_payload)
            if isinstance(accepted_payload, dict)
            else accepted_payload
        )
        if not isinstance(accepted_edits, list):
            raise SystemExit("Accepted edits JSON must be a list or an object with suggested_edits.")
        apply_edits(args.resume, args.out, accepted_edits)
        print(f"Wrote tailored resume: {args.out}")
    else:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.resume, args.out)
        print(f"No accepted edits provided; copied original resume to: {args.out}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
