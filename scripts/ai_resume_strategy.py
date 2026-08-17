#!/usr/bin/env python3
"""Build an auditable, truth-bound resume recommendation strategy."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any, Iterable


SKILL_ALIASES: dict[str, tuple[str, ...]] = {
    "data engineering": ("data engineer", "data engineering"),
    "data pipelines": ("data pipeline", "data pipelines", "etl pipeline", "etl pipelines", "elt pipeline"),
    "data modeling": ("data model", "data models", "data modeling", "data modelling", "schema design"),
    "data quality": ("data quality", "data validation", "reconciliation"),
    "data warehouse": ("data warehouse", "data warehousing", "warehouse"),
    "batch processing": ("batch processing", "batch pipeline", "batch pipelines", "scheduled processing"),
    "stream processing": ("stream processing", "streaming", "real-time pipeline", "real time pipeline"),
    "distributed systems": ("distributed system", "distributed systems"),
    "stakeholder partnership": ("stakeholder", "business partner", "cross-functional", "cross functional"),
    "requirements definition": ("requirements", "acceptance criteria", "business rules"),
    "testing": ("unit test", "unit testing", "test case", "testing", "uat"),
    "documentation": ("documentation", "documenting", "technical documentation"),
    "sla operations": ("sla", "service level", "production support", "operational support"),
    "ci/cd": ("ci/cd", "cicd", "continuous integration", "continuous delivery"),
    "governance": ("governance", "controls", "audit"),
    "analytics": ("analytics", "business intelligence", "reporting", "dashboard"),
    "sql": ("sql", "sql server", "postgres", "postgresql"),
    "python": ("python",),
    "spark": ("spark", "pyspark"),
    "airflow": ("airflow", "apache airflow"),
    "dbt": ("dbt",),
    "snowflake": ("snowflake",),
    "kafka": ("kafka", "apache kafka"),
    "aws": ("aws", "amazon web services"),
    "gcp": ("gcp", "google cloud", "google cloud platform"),
    "azure": ("azure", "microsoft azure"),
    "tableau": ("tableau",),
    "git": ("git", "gitlab", "github"),
    "ssis": ("ssis", "sql server integration services"),
    "c#": ("c#", "c sharp"),
    "machine learning": ("machine learning", "ml model", "ml models"),
    "llm workflows": ("llm", "large language model", "generative ai"),
}

TRANSFERABLE_EVIDENCE: dict[str, tuple[str, ...]] = {
    "dbt": ("sql", "etl", "elt", "data pipelines", "data modeling", "data quality", "data warehouse"),
    "snowflake": ("sql", "etl", "elt", "data pipelines", "data modeling", "data quality", "data warehouse"),
    "airflow": ("batch processing", "orchestration", "scheduled processing", "dependency-based sequencing"),
    "kafka": ("data pipelines", "stream processing"),
}

CONCRETE_TOOL_SKILLS = {
    "airflow",
    "aws",
    "azure",
    "c#",
    "dbt",
    "gcp",
    "git",
    "kafka",
    "python",
    "snowflake",
    "spark",
    "sql",
    "ssis",
    "tableau",
}

NEGATIVE_EVIDENCE_CUES = (
    "do not claim",
    "no professional",
    "no hands-on",
    "not hands-on",
    "not treat it as",
    "do not treat",
    "must not",
    "unless confirmed",
    "unless that experience is confirmed",
    "without confirmation",
    "transferable",
    "workflow equivalent",
)

MUST_SECTION_CUES = (
    "requirements",
    "minimum qualifications",
    "basic qualifications",
    "what you bring",
    "what you'll bring",
    "what you will bring",
    "who you are",
)
NICE_SECTION_CUES = ("preferred", "nice to have", "bonus", "ideally")
MUST_LINE_CUES = ("must", "required", "requirement", "minimum", "at least", "years of")
NICE_LINE_CUES = ("preferred", "nice to have", "bonus", "a plus", "ideally")


def canonical_token(token: str) -> str:
    token = token.lower()
    if len(token) <= 3 or token.endswith(("ss", "us")):
        return token
    if token.endswith("ies") and len(token) > 4:
        return token[:-3] + "y"
    if token.endswith("ses") and len(token) > 5:
        return token[:-2]
    if token.endswith("s"):
        return token[:-1]
    return token


def canonical_text(text: str) -> str:
    words = re.findall(r"[a-z0-9+#]+", text.lower().replace("&", " and "))
    return " ".join(canonical_token(word) for word in words)


def phrase_present(text: str, phrase: str) -> bool:
    candidate = canonical_text(text)
    target = canonical_text(phrase)
    return bool(target and re.search(rf"(?<!\w){re.escape(target)}(?!\w)", candidate))


def _lines(text: str) -> list[str]:
    return [
        re.sub(r"^[\s*\-\u2022\u25cf\u25e6]+", "", line).strip()
        for line in text.splitlines()
        if line.strip()
    ]


def _priority_by_line(job_text: str) -> dict[str, str]:
    priorities: dict[str, str] = {}
    current = "core"
    for line in _lines(job_text):
        lower = line.lower().rstrip(":")
        if len(line) <= 80 and any(cue in lower for cue in MUST_SECTION_CUES):
            current = "must"
            continue
        if len(line) <= 80 and any(cue in lower for cue in NICE_SECTION_CUES):
            current = "nice"
            continue
        line_priority = current
        if any(cue in lower for cue in NICE_LINE_CUES):
            line_priority = "nice"
        elif any(cue in lower for cue in MUST_LINE_CUES):
            line_priority = "must"
        for skill, aliases in SKILL_ALIASES.items():
            if any(phrase_present(line, alias) for alias in aliases):
                previous = priorities.get(skill)
                if previous != "must" and (line_priority == "must" or previous is None):
                    priorities[skill] = line_priority
                elif previous == "nice" and line_priority == "core":
                    priorities[skill] = "core"
    return priorities


def _evidence_lines(resume_paragraphs: Iterable[Any], profile_facts: str) -> list[dict[str, str]]:
    evidence: list[dict[str, str]] = []
    for paragraph in resume_paragraphs:
        text = str(getattr(paragraph, "text", paragraph)).strip()
        if text:
            evidence.append({"source": "resume", "text": text})
    for line in _lines(profile_facts):
        if line:
            evidence.append({"source": "profile", "text": line})
    return evidence


def _direct_evidence(skill: str, records: list[dict[str, str]]) -> dict[str, str] | None:
    aliases = SKILL_ALIASES[skill]
    for source in ("resume", "profile"):
        for record in records:
            if record["source"] != source:
                continue
            lower = record["text"].lower()
            if any(cue in lower for cue in NEGATIVE_EVIDENCE_CUES):
                continue
            if source == "profile" and skill in CONCRETE_TOOL_SKILLS:
                positive_cues = ("confirmed", "hands-on", "professional experience", "used ", "built ", "developed ")
                if not any(cue in lower for cue in positive_cues):
                    continue
            if any(phrase_present(record["text"], alias) for alias in aliases):
                return record
    return None


def _transferable_evidence(skill: str, records: list[dict[str, str]]) -> dict[str, str] | None:
    workflows = TRANSFERABLE_EVIDENCE.get(skill, ())
    if not workflows:
        return None
    for record in records:
        if any(phrase_present(record["text"], workflow) for workflow in workflows):
            return record
    return None


def _shorten(text: str, limit: int = 220) -> str:
    return text if len(text) <= limit else text[: limit - 3].rstrip() + "..."


def audit_docx_parser(path: Path | None) -> dict[str, Any]:
    if path is None or not path.exists():
        return {"risk": "unknown", "checks": [], "recommendation": "Run parser audit on the selected DOCX."}
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return {"risk": "unknown", "checks": [], "recommendation": "Install python-docx to audit parsing risk."}

    document = Document(path)
    checks: list[dict[str, Any]] = []
    table_count = len(document.tables)
    checks.append({"name": "layout_tables", "value": table_count, "pass": table_count == 0})
    header_text = " ".join(p.text.strip() for section in document.sections for p in section.header.paragraphs).strip()
    footer_text = " ".join(p.text.strip() for section in document.sections for p in section.footer.paragraphs).strip()
    checks.append({"name": "header_footer_text", "value": bool(header_text or footer_text), "pass": not bool(header_text or footer_text)})
    multi_column = False
    for section in document.sections:
        columns = section._sectPr.xpath("./w:cols")
        if columns:
            count = columns[0].get(qn("w:num"))
            multi_column = multi_column or (count is not None and int(count) > 1)
    checks.append({"name": "multiple_columns", "value": multi_column, "pass": not multi_column})
    image_count = len(document.inline_shapes)
    checks.append({"name": "inline_images", "value": image_count, "pass": image_count == 0})
    text_boxes = "w:txbxContent" in document.element.body.xml
    checks.append({"name": "text_boxes", "value": text_boxes, "pass": not text_boxes})
    size_bytes = path.stat().st_size
    checks.append({"name": "file_size_under_2_5mb", "value": size_bytes, "pass": size_bytes <= 2_500_000})

    failed = [check["name"] for check in checks if not check["pass"]]
    risk = "low" if not failed else "high" if any(name in failed for name in ("layout_tables", "multiple_columns", "text_boxes")) else "medium"
    return {
        "risk": risk,
        "checks": checks,
        "failed_checks": failed,
        "recommendation": (
            "Use a single-column DOCX with standard section headings and contact details in the body."
            if failed
            else "The selected DOCX follows conservative parser-safe structure."
        ),
    }


def build_selection_report(
    job_text: str,
    resume_paragraphs: Iterable[Any],
    profile_facts: str,
    *,
    resume_path: Path | None = None,
) -> dict[str, Any]:
    paragraphs = list(resume_paragraphs)
    records = _evidence_lines(paragraphs, profile_facts)
    priorities = _priority_by_line(job_text)
    criteria: list[dict[str, Any]] = []
    for skill, priority in priorities.items():
        direct = _direct_evidence(skill, records)
        transferable = None if direct else _transferable_evidence(skill, records)
        if direct:
            status, confidence, evidence = "supported", "high", direct
        elif transferable:
            status, confidence, evidence = "transferable", "medium", transferable
        else:
            status, confidence, evidence = "unsupported", "low", None
        criteria.append(
            {
                "criterion": skill,
                "priority": priority,
                "status": status,
                "confidence": confidence,
                "job_terms": [alias for alias in SKILL_ALIASES[skill] if phrase_present(job_text, alias)],
                "evidence_source": evidence["source"] if evidence else None,
                "evidence": _shorten(evidence["text"]) if evidence else None,
            }
        )

    priority_order = {"must": 0, "core": 1, "nice": 2}
    status_order = {"supported": 0, "transferable": 1, "unsupported": 2}
    criteria.sort(key=lambda item: (priority_order[item["priority"]], status_order[item["status"]], item["criterion"]))
    weights = {"must": 3.0, "core": 2.0, "nice": 1.0}
    values = {"supported": 1.0, "transferable": 0.65, "unsupported": 0.0}
    denominator = sum(weights[item["priority"]] for item in criteria)
    numerator = sum(weights[item["priority"]] * values[item["status"]] for item in criteria)
    coverage = round(100 * numerator / denominator) if denominator else 0
    prominent = [
        item["criterion"]
        for item in criteria
        if item["status"] == "supported" and item["priority"] in {"must", "core"}
    ][:8]
    return {
        "strategy_version": 1,
        "coverage_score": coverage,
        "criteria": criteria,
        "prominence_plan": {
            "summary": prominent[:4],
            "skills_section": prominent,
            "experience": [item["criterion"] for item in criteria if item["status"] == "supported"][:10],
        },
        "parser_audit": audit_docx_parser(resume_path),
        "guardrails": [
            "Use exact job terminology only when supported by resume or private profile evidence.",
            "Put the strongest must-have evidence in the summary, skills section, and recent experience.",
            "Do not add invisible text, keyword lists without context, prompt injection, or fabricated experience.",
            "Keep unsupported and transferable tool gaps explicit; workflow similarity is not hands-on tool experience.",
        ],
        "model_scope": (
            "Optimizes observable parsing and criterion coverage. It does not reproduce or guarantee any employer's proprietary ranking."
        ),
    }


def validate_edit_against_strategy(edit: dict[str, Any], report: dict[str, Any]) -> list[str]:
    original = str(edit.get("original", ""))
    suggested = str(edit.get("suggested", ""))
    failures: list[str] = []
    for criterion in report.get("criteria", []):
        if criterion.get("status") == "supported":
            continue
        terms = criterion.get("job_terms") or [criterion.get("criterion", "")]
        for term in terms:
            if term and phrase_present(suggested, term) and not phrase_present(original, term):
                failures.append(
                    f"New {criterion.get('status')} criterion claim is not allowed: {criterion.get('criterion')}."
                )
                break
    for criterion in report.get("criteria", []):
        term = str(criterion.get("criterion", ""))
        if not term:
            continue
        before = len(re.findall(re.escape(canonical_text(term)), canonical_text(original)))
        after = len(re.findall(re.escape(canonical_text(term)), canonical_text(suggested)))
        if after > max(2, before + 1):
            failures.append(f"Possible keyword stuffing for criterion: {term}.")
    return failures


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build an evidence-backed AI resume selection report.")
    parser.add_argument("--job", type=Path, required=True)
    parser.add_argument("--resume", type=Path, default=Path("resumes/master_ats.docx"))
    parser.add_argument("--profile", type=Path, default=Path("profile/facts.md"))
    parser.add_argument("--out", type=Path, required=True)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    try:
        from tailor import load_resume_paragraphs
    except ImportError:  # pragma: no cover
        from scripts.tailor import load_resume_paragraphs
    if not args.job.exists():
        raise SystemExit(f"Job description not found: {args.job}")
    if not args.resume.exists():
        raise SystemExit(f"Resume not found: {args.resume}")
    job_text = args.job.read_text(encoding="utf-8")
    profile_facts = args.profile.read_text(encoding="utf-8") if args.profile.exists() else ""
    report = build_selection_report(
        job_text,
        load_resume_paragraphs(args.resume),
        profile_facts,
        resume_path=args.resume,
    )
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(f"Wrote AI selection report: {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
