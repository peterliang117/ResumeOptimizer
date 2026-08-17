#!/usr/bin/env python3
"""Create a single-column, parser-safe resume from the local master DOCX."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any


def _set_cell_free_styles(document: Any) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.25)
    normal.font.color.rgb = RGBColor(20, 24, 28)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0.8)
    normal.paragraph_format.line_spacing = 1.0

    styles = {
        "Resume Name": (15.5, True, 0, 0),
        "Resume Contact": (8.0, False, 0, 3),
        "Resume Section": (9.0, True, 3, 1),
        "Resume Entry": (8.4, True, 1, 0),
        "Resume Detail": (8.1, False, 0, 0.5),
    }
    for name, (size, bold, before, after) in styles.items():
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(8.05)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(0.6)
    bullet.paragraph_format.line_spacing = 1.0
    bullet.paragraph_format.left_indent = Inches(0.16)
    bullet.paragraph_format.first_line_indent = Inches(-0.12)


def _add_bottom_border(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2C5F73")
    borders.append(bottom)
    p_pr.append(borders)


def _add_heading(document: Any, text: str) -> None:
    paragraph = document.add_paragraph(style="Resume Section")
    paragraph.add_run(text.upper())
    _add_bottom_border(paragraph)


def _add_entry(document: Any, title: str, meta: list[str]) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches

    paragraph = document.add_paragraph(style="Resume Entry")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run(title.strip())
    clean_meta = [item.strip() for item in meta if item.strip()]
    if len(clean_meta) >= 2 and clean_meta[-2].endswith(("-", "–", "—")):
        clean_meta[-2:] = [f"{clean_meta[-2]} {clean_meta[-1]}"]
    if clean_meta:
        paragraph.add_run("\t" + " | ".join(clean_meta))


def _table_rows(table: Any) -> list[tuple[list[str], list[str]]]:
    def cell_lines(cell: Any) -> list[str]:
        # Some source-resume locations are stored in Word content controls.
        # python-docx excludes that text from ``paragraph.text``, so read the
        # text nodes from each direct cell paragraph instead.
        lines = []
        for paragraph in cell._tc.xpath(".//w:p[not(ancestor::w:p)]"):
            text = "".join(node.text or "" for node in paragraph.xpath(".//w:t")).strip()
            if text:
                lines.append(text)
        return lines

    rows: list[tuple[list[str], list[str]]] = []
    for row in table.rows:
        left = cell_lines(row.cells[0])
        right = []
        if len(row.cells) > 1:
            right = cell_lines(row.cells[1])
        rows.append((left, right))
    return rows


def build_ats_resume(source: Path, output: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    source_doc = Document(source)
    if len(source_doc.paragraphs) < 2 or len(source_doc.tables) < 4:
        raise SystemExit("The source resume does not match the expected section structure.")

    document = Document()
    _set_cell_free_styles(document)

    name = document.add_paragraph(style="Resume Name")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.add_run(source_doc.paragraphs[0].text.strip())
    contact = document.add_paragraph(style="Resume Contact")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(source_doc.paragraphs[1].text.strip())

    _add_heading(document, "Summary")
    summary = _table_rows(source_doc.tables[0])[0][0]
    for line in summary:
        document.add_paragraph(line, style="Resume Detail")

    _add_heading(document, "Experience")
    for left, right in _table_rows(source_doc.tables[2]):
        if not left:
            continue
        _add_entry(document, left[0], right)
        for bullet in left[1:]:
            document.add_paragraph(bullet, style="List Bullet")

    _add_heading(document, "Technical Capabilities")
    for line in _table_rows(source_doc.tables[3])[0][0]:
        document.add_paragraph(line, style="Resume Detail")

    _add_heading(document, "Education")
    for left, right in _table_rows(source_doc.tables[1]):
        if not left:
            continue
        _add_entry(document, left[0], right)
        for detail in left[1:]:
            document.add_paragraph(detail, style="Resume Detail")

    if len(source_doc.tables) > 4:
        additional = _table_rows(source_doc.tables[4])[0][0]
        if additional:
            _add_heading(document, "Additional Information")
            for line in additional:
                document.add_paragraph(line, style="Resume Detail")

    candidate_name = source_doc.paragraphs[0].text.strip() or "Candidate"
    document.core_properties.title = f"{candidate_name} Resume"
    document.core_properties.subject = "ATS-safe professional resume"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def update_manifest(path: Path, resume_path: Path) -> None:
    if not path.exists():
        return
    payload = json.loads(path.read_text(encoding="utf-8"))
    for variant in payload.get("variants", {}).values():
        if isinstance(variant, dict):
            variant["resume_file"] = str(resume_path)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a parser-safe single-column resume.")
    parser.add_argument("--source", type=Path, default=Path("resumes/master.docx"))
    parser.add_argument("--out", type=Path, default=Path("resumes/master_ats.docx"))
    parser.add_argument("--manifest", type=Path, default=Path("profile/resume_variants.private.json"))
    parser.add_argument("--update-manifest", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if not args.source.exists():
        raise SystemExit(f"Resume not found: {args.source}")
    build_ats_resume(args.source, args.out)
    if args.update_manifest:
        update_manifest(args.manifest, args.out)
    print(f"Wrote ATS-safe resume: {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
